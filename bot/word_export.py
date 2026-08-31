from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

def export_cp_to_word(cp_data: dict) -> BytesIO:
    """Экспортирует КП в Word"""
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Информация о поставщике
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"Поставщик: ").bold = True
    p.add_run(cp_data.get('supplier_name', ''))
    
    p = doc.add_paragraph()
    p.add_run(f"ИНН: ").bold = True
    p.add_run(cp_data.get('supplier_inn', ''))
    
    p = doc.add_paragraph()
    p.add_run(f"Контактное лицо: ").bold = True
    p.add_run(cp_data.get('contact', ''))
    
    doc.add_paragraph()
    
    # Информация о тендере
    p = doc.add_paragraph()
    p.add_run(f"Тендер: ").bold = True
    p.add_run(cp_data.get('tender_name', ''))
    
    if cp_data.get('tender_number'):
        p = doc.add_paragraph()
        p.add_run(f"Номер тендера: ").bold = True
        p.add_run(cp_data.get('tender_number', ''))
    
    doc.add_paragraph()
    
    # Таблица позиций
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Light Grid Accent 1'
    
    headers = ['№', 'Наименование', 'Производитель', 'Модель', 'Кол-во', 'Ед.', 'Цена, ₽', 'Сумма, ₽']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    # Данные
    for item in cp_data.get('items', []):
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.get('position', ''))
        row_cells[1].text = item.get('name', '')
        row_cells[2].text = item.get('manufacturer', '')
        row_cells[3].text = item.get('model', '')
        row_cells[4].text = str(item.get('quantity', 1))
        row_cells[5].text = item.get('unit', 'шт')
        row_cells[6].text = f"{item.get('price', 0):,.2f}"
        row_cells[7].text = f"{item.get('sum', 0):,.2f}"
    
    # Итого
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"Итого без НДС: ").bold = True
    p.add_run(f"{cp_data.get('total', 0):,.2f} ₽")
    
    p = doc.add_paragraph()
    p.add_run(f"НДС 20%: ").bold = True
    p.add_run(f"{cp_data.get('vat', 0):,.2f} ₽")
    
    p = doc.add_paragraph()
    run = p.add_run(f"Всего с НДС: ")
    run.bold = True
    run.font.size = Pt(14)
    run = p.add_run(f"{cp_data.get('total_with_vat', 0):,.2f} ₽")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 255)
    
    # Условия
    doc.add_paragraph()
    doc.add_heading('Условия поставки:', level=2)
    
    doc.add_paragraph(f"• Срок поставки: {cp_data.get('delivery_days', '')} дней (до {cp_data.get('delivery_date', '')})")
    doc.add_paragraph(f"• Условия оплаты: {cp_data.get('payment_terms', '')}")
    doc.add_paragraph(f"• Гарантия: {cp_data.get('warranty', '')}")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"Дата формирования: {cp_data.get('generated_at', '')}").italic = True
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
