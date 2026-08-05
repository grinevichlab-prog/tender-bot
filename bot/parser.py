"""
Модуль извлечения текста из файлов разных форматов.
Поддержка .doc/.xls через LibreOffice на Linux.
"""

import os
import sys
import subprocess
import tempfile
import zipfile
from pathlib import Path

import PyPDF2
import pdfplumber
import docx
import openpyxl
from PIL import Image
import pytesseract

if sys.platform == "win32":
    try:
        import win32com.client
    except ImportError:
        win32com = None
else:
    win32com = None


def _convert_via_libreoffice(file_path: str, target_ext: str) -> str:
    """
    Конвертирует файл (.doc -> docx, .xls -> xlsx) через LibreOffice (headless).
    Каждый вызов получает свой изолированный профиль, чтобы избежать
    зависания из-за файла-замка при параллельных конвертациях.
    """
    output_dir = tempfile.mkdtemp()
    profile_dir = tempfile.mkdtemp()
    profile_uri = f"file://{profile_dir}"

    try:
        subprocess.run(
            [
                "libreoffice",
                f"-env:UserInstallation={profile_uri}",
                "--headless", "--norestore", "--nologo",
                "--convert-to", target_ext,
                "--outdir", output_dir, file_path,
            ],
            check=True,
            timeout=60,
        )
    except subprocess.TimeoutExpired as e:
        raise RuntimeError(f"LibreOffice завис при конвертации {file_path}: {e}")

    for f in os.listdir(output_dir):
        if f.endswith(f".{target_ext}"):
            return os.path.join(output_dir, f)
    raise RuntimeError(f"Не удалось найти сконвертированный .{target_ext}")


def _extract_text_from_docx(file_path: str) -> str:
    doc = docx.Document(file_path)
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def _extract_text_from_doc(file_path: str) -> str:
    if sys.platform == "win32" and win32com:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(file_path)
            text = doc.Content.Text
            doc.Close()
            return text
        finally:
            word.Quit()
    else:
        docx_path = _convert_via_libreoffice(file_path, "docx")
        try:
            return _extract_text_from_docx(docx_path)
        finally:
            try:
                os.remove(docx_path)
            except Exception:
                pass


def _extract_text_from_pdf(file_path: str) -> str:
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception:
        pass

    if not text.strip():
        try:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception:
            pass
    return text


def _extract_text_from_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _extract_text_from_xlsx(file_path: str) -> str:
    wb = openpyxl.load_workbook(file_path, data_only=True)
    all_text = []
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join(
                str(cell) if cell is not None else "" for cell in row
            )
            if row_text.strip():
                all_text.append(row_text)
    return "\n".join(all_text)


def _extract_text_from_xls(file_path: str) -> str:
    if sys.platform == "win32" and win32com:
        excel = win32com.client.Dispatch("Excel.Application")
        excel.Visible = False
        try:
            wb = excel.Workbooks.Open(file_path)
            lines = []
            for sheet in wb.Sheets:
                used = sheet.UsedRange
                for row in used.Rows:
                    values = [str(c.Value) for c in row.Cells if c.Value is not None]
                    if values:
                        lines.append(" | ".join(values))
            wb.Close(False)
            return "\n".join(lines)
        finally:
            excel.Quit()
    else:
        xlsx_path = _convert_via_libreoffice(file_path, "xlsx")
        try:
            return _extract_text_from_xlsx(xlsx_path)
        finally:
            try:
                os.remove(xlsx_path)
            except Exception:
                pass


def _ocr_image(file_path: str) -> str:
    img = Image.open(file_path)
    return pytesseract.image_to_string(img, lang="rus+eng")


def extract_texts_from_zip(file_path: str) -> list[dict]:
    results = []
    with tempfile.TemporaryDirectory() as tmpdir:
        with zipfile.ZipFile(file_path, "r") as zf:
            zf.extractall(tmpdir)

        for root, dirs, files in os.walk(tmpdir):
            for fname in files:
                full_path = os.path.join(root, fname)
                ext = Path(fname).suffix.lower()
                try:
                    text = extract_text(full_path, ext)
                    if text and text.strip():
                        results.append({"name": fname, "text": text})
                except Exception as e:
                    print(f"Ошибка при обработке файла в архиве {fname}: {e}", flush=True)
    return results


def extract_text(file_path: str, file_ext: str) -> str:
    ext = file_ext.lower()
    print(f"[extract_text] начинаю обработку {file_path} ({ext})", flush=True)

    if ext == ".pdf":
        result = _extract_text_from_pdf(file_path)
    elif ext == ".docx":
        result = _extract_text_from_docx(file_path)
    elif ext == ".doc":
        result = _extract_text_from_doc(file_path)
    elif ext in (".xlsx", ".xlsm"):
        result = _extract_text_from_xlsx(file_path)
    elif ext == ".xls":
        result = _extract_text_from_xls(file_path)
    elif ext == ".txt":
        result = _extract_text_from_txt(file_path)
    elif ext in (".png", ".jpg", ".jpeg"):
        result = _ocr_image(file_path)
    elif ext == ".zip":
        items = extract_texts_from_zip(file_path)
        combined = [f"=== {i['name']} ===\n{i['text']}" for i in items]
        result = "\n\n".join(combined)
    else:
        raise ValueError(f"Неподдерживаемый формат файла: {ext}")

    print(f"[extract_text] готово, символов: {len(result)}", flush=True)
    return result
